# -*- coding: utf-8 -*-
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

from openpyxl import load_workbook
import pandas as pd
from docx import Document
import docx
import numpy as np

# data_only = True로 해줘야 수식이 아닌 값으로 받아온다.
load_wb = load_workbook("C:\Users\JeongHwanSeock\PycharmProjects\PDF\SA12.xlsx")
# 시트이름으로 불러오기
load_ws = load_wb['Index']
all_value = []
for row in load_ws.rows:
    row_value = []
    for cell in row:
        row_value.append(cell.value)
    all_value.append(row_value)
all_value = pd.DataFrame(all_value[1:]).sort_values(by=3)
p_title = all_value[1].dropna(axis=0).reset_index(drop=True)  #테이블 제목
r_title = all_value[2].dropna(axis=0).reset_index(drop=True)  # 결과 검토
result_summary = str(load_ws['A2'].value)
load_ws2 = load_wb[result_summary]

all_values = []
for row in load_ws2.rows:
    row_value = []
    for cell in row:
        row_value.append(cell.value)
    all_values.append(row_value)
column = all_values[0]
data = all_values[1:]
df = pd.DataFrame(data=data, columns=column)
df2 = df[['Power Level (%)', 'Iteration','PF Target','PF Actual 1', 'PF Actual 2', 'PF Actual 3']]

#사용하기 위한 변수 선언
document = Document()
# title = input('시험 제목을 입력하시오: ')
# description = input('시험 설명을 입력하시오: ')
title = 'SPF(Specified Powr Factor) 기능'.decode('utf-8')
description = "- 측정된 역률(Power Factor) 값과 설정한 역률 값이 차이가 제조사가 명시한 정확도 (Manufacturer's Stated Accuracy)내에 있는지 여부로 판단. " \
              "피시험 인버터인 STP12000TL-US-10의 역률 정확도는 0.01이고 설정 가능 범위는 Minimum Capacitive Power Factor는 0.8, " \
              "Minimum Inductive (Underexcited) Power Factor는 –0.8임.".decode('utf-8')

#제목을 위한 add_heading('제목',0)함수사용
document.add_heading(title, 0)
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER # 중앙정렬

document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
document.add_paragraph(description)
document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
num = df2['PF Target'].unique()
num = np.sort(num[1:])
num2 = 1

# 테이블 작성
for i in range(len(p_title)):
    a = p_title[i]
    mer_title = str(num2)  + '.  ' + a
    document.add_paragraph(mer_title)
    table = document.add_table(rows=1, cols=6, style='Light Shading')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '출력 (%)'.decode('utf-8')
    hdr_cells[1].text = '반복횟수'.decode('utf-8')
    hdr_cells[2].text = '목표 역률'.decode('utf-8')
    hdr_cells[3].text = '실제 역률 (A)'.decode('utf-8')
    hdr_cells[4].text = '실제 역률 (B)'.decode('utf-8')
    hdr_cells[5].text = '실제 역률 (C)'.decode('utf-8')
    for a,b,c,d,e,f in df2[df2['PF Target']==num[i]].values.tolist():
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)
        row_cells[3].text = str(d)
        row_cells[4].text = str(e)
        row_cells[5].text = str(f)
    num2 +=1
    try:
        temp = r_title[i]
        mer_title2 = '* 결과검토: ' + temp
        document.add_paragraph(mer_title2)
    except KeyError:
        pass
document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'))



# docx파일을 생성을 위한 save('파일명')
document.save('demo.docx')