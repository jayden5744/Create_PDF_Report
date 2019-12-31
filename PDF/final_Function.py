# -*- coding: utf-8 -*-
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

from openpyxl import load_workbook
import pandas as pd
from docx import Document
import docx
from docx.shared import Inches, RGBColor
import numpy as np
import os
import comtypes.client
import matplotlib.pyplot as plt
from docx.oxml.ns import qn

# 스타일 선언
def style(document):
    style_1 = document.styles.add_style('Heading_1', docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
    style_1.base_style = document.styles['Heading 1']
    style_1.font.name = 'HY견고딕'.decode('utf-8')
    style_1._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY견고딕'.decode('utf-8')) # 한글 폰트를 따로 설정해 준다
    style_1.font.size = docx.shared.Pt(15)
    style_1.font.bold = True
    style_1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    return style_1


def SA12(filename):
    print('-----------------SA12 연구파일 pdf변환을 시작합니다.---------------------')
    # data_only = True로 해줘야 수식이 아닌 값으로 받아온다.
    print('-----------------Excel File을 성공적으로 불러왔습니다.---------------------')
    load_wb = load_workbook("C:\Users\JeongHwanSeock\PycharmProjects\PDF\\"+filename+".xlsx")
    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:]).sort_values(by=3)
    p_title = all_value[1].dropna(axis=0).reset_index(drop=True)  # 테이블 제목
    r_title = all_value[2].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    result_summary = str(load_ws['A2'].value)
    load_ws2 = load_wb[result_summary]

    all_values = []
    for row in load_ws2.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_values.append(row_value)
    column = all_values[0]
    data = all_values[1:]
    df = pd.DataFrame(data=data, columns=column)
    df2 = df[['Power Level (%)', 'Iteration', 'PF Target', 'PF Actual 1', 'PF Actual 2', 'PF Actual 3']]

    # 사용하기 위한 변수 선언
    document = Document()
    # title = input('시험 제목을 입력하시오: ')
    # description = input('시험 설명을 입력하시오: ')
    title = 'SPF(Specified Powr Factor) 기능'.decode('utf-8')
    description = "- 측정된 역률(Power Factor) 값과 설정한 역률 값이 차이가 제조사가 명시한 정확도 (Manufacturer's Stated Accuracy)내에 있는지 여부로 판단. " \
                  "피시험 인버터인 STP12000TL-US-10의 역률 정확도는 0.01이고 설정 가능 범위는 Minimum Capacitive Power Factor는 0.8, " \
                  "Minimum Inductive (Underexcited) Power Factor는 –0.8임.".decode('utf-8')

    # 제목
    style_1 = style(document)
    document.add_paragraph(title, style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험 설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(description)
    # 기능시험 결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    num = df2['PF Target'].unique()
    num = np.sort(num[1:])

    # 테이블 작성
    for i in range(len(p_title)):
        mer_title = p_title[i]
        document.add_paragraph(mer_title, style='ListNumber')
        table = document.add_table(rows=1, cols=6, style='Light Shading')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '출력 (%)'.decode('utf-8')
        hdr_cells[1].text = '반복횟수'.decode('utf-8')
        hdr_cells[2].text = '목표 역률'.decode('utf-8')
        hdr_cells[3].text = '실제 역률 (A)'.decode('utf-8')
        hdr_cells[4].text = '실제 역률 (B)'.decode('utf-8')
        hdr_cells[5].text = '실제 역률 (C)'.decode('utf-8')
        for a, b, c, d, e, f in df2[df2['PF Target'] == num[i]].values.tolist():
            row_cells = table.add_row().cells
            row_cells[0].text = str(a)
            row_cells[1].text = str(b)
            row_cells[2].text = str(c)
            row_cells[3].text = str(d)
            row_cells[4].text = str(e)
            row_cells[5].text = str(f)
        try:
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2)
        except KeyError:
            pass
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='List Bullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'))

    # docx파일을 생성을 위한 save('파일명')
    document.save('demo1.docx')
    print('-----------------Docs File을 성공적으로 불러왔습니다.---------------------')

    wdFormatPDF = 17
    # 파일 경로 절대경로로
    in_file = os.path.abspath('demo1.docx')
    out_file = os.path.abspath(filename)
    # word형식의 파일을 열기
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    # PDF형식으로 저장
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print('-----------------PDF File을 성공적으로 만들었습니다.---------------------')

def SA13(filename):
    print('-----------------SA13 연구파일 pdf변환을 시작합니다.---------------------')
    print('-----------------Excel File을 성공적으로 불러왔습니다.---------------------')
    filename = 'SA13'
    # data_only = True로 해줘야 수식이 아닌 값으로 받아온다.
    load_wb = load_workbook("C:\Users\JeongHwanSeock\PycharmProjects\PDF\\" + filename + ".xlsx")
    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    r_title = all_value[2].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    result_summary = str(load_ws['A2'].value)
    load_ws2 = load_wb[result_summary]

    all_values = []
    for row in load_ws2.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_values.append(row_value)
    column = all_values[0]
    data = all_values[1:]
    df = pd.DataFrame(data=data, columns=column)
    # 그림 만들기
    index_down = df[df['Dataset File'].str.contains('down') == True].index
    index_up = df[df['Dataset File'].str.contains('up') == True].index
    for i in range(len(index_down)):
        fig = plt.figure()
        plt.rcParams["figure.figsize"] = (10, 6)
        plt.rcParams['axes.grid'] = True

        img_title = df['Dataset File'][index_down[i]]
        if (i == 0):
            VV_1_1000 = df[0:index_down[i]]
        else:
            VV_1_1000 = df[index_up[i - 1] + 1:index_down[i]]
        ax = fig.add_subplot(1, 1, 1)
        ax.plot(VV_1_1000['Average Voltage (pu)'], VV_1_1000['Var Actual 1'] / 4000, linestyle='', marker='o',
                color='blue', label='Power')
        ax.plot(VV_1_1000['Average Voltage (pu)'], VV_1_1000['Var Target 1'] / 4000, color='black', label='VV curve')
        ax.plot(VV_1_1000['Average Voltage (pu)'], VV_1_1000['Var Min Allowed 1'] / 4000, linestyle=':', color='red',
                label='VV pass/fail band')
        ax.plot(VV_1_1000['Average Voltage (pu)'], VV_1_1000['Var Max Allowed 1'] / 4000, linestyle=':', color='red')
        ax.set_title('Volt-Var Function1')
        ax.set_xlabel('Grid Voltage(% nominal)')
        ax.set_ylabel('Reactive Power(% nameplate)')
        ax.set_xticks([0.9, 0.95, 1, 1.05, 1.1])
        ax.set_xticklabels(['90', '95', '100', '105', '110'])
        ax.set_yticks([-1.5, -1, -0.5, 0, 0.5, 1.0, 1.5])
        ax.set_yticklabels(['-150', '-100', '-50', '0', '50', '100', '150'])

        plt.savefig('img/' + img_title + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    # title = input('시험 제목을 입력하시오: ')
    # description = input('시험 설명을 입력하시오: ')
    title = 'Volt-Var 기능 (Most Aggressive Curve)'.decode('utf-8')
    description = "사용자입력".decode('utf-8')

    # 제목
    style_1 = style(document)  # 스타일 설정
    document.add_paragraph(title, style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(description)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'))

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(index_down)):
        img_title = df['Dataset File'][index_down[i]]
        mer_title = str(p_title[i])
        document.add_paragraph(mer_title.decode('utf-8'), style='ListNumber')
        document.add_picture('img/' + str(img_title) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + str(mer_title) + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2)
        except KeyError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'))

    # docx파일을 생성을 위한 save('파일명')
    document.save('demo2.docx')
    print('-----------------Docs File을 성공적으로 불러왔습니다.---------------------')

    wdFormatPDF = 17

    # 파일 경로 절대경로로
    in_file = os.path.abspath('demo2.docx')
    out_file = os.path.abspath(filename)
    # word형식의 파일을 열기
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    # PDF형식으로 저장
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print('-----------------PDF File을 성공적으로 만들었습니다.---------------------')
SA12('SA12')
SA13('SA13')
