# -*- coding: utf-8 -*-
import os
import docx
import xlrd
import math
import argparse
import openpyxl
import datetime
import numpy as np
import pandas as pd
import excel2pdf_gui
import comtypes.client
from docx import Document
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
from openpyxl import load_workbook
from docx.shared import Inches, RGBColor
from comtypes import COMError
import sys

reload(sys)
sys.setdefaultencoding('utf-8')
plt.rcParams['axes.grid'] = True

# 그래프에서 한글을 쓰고 싶을 경우
# import matplotlib
# import matplotlib.font_manager as fm
# font_location = 'C:\\WINDOWS\\Fonts\\NanumBarunGothic.ttf'  # 폰트의 경로
# fon_name =fm.FontProperties(fname=font_location).get_name()
# matplotlib.rc('font', family=fon_name)


class TkErrorCatcher:
    '''
    In some cases tkinter will only print the traceback.
    Enables the program to catch tkinter errors normally

    To use
    import tkinter
    tkinter.CallWrapper = TkErrorCatcher
    '''

    def __init__(self, func, subst, widget):
        self.func = func
        self.subst = subst
        self.widget = widget

    def __call__(self, *args):
        try:
            if self.subst:
                args = self.subst(*args)
            return self.func(*args)
        except SystemExit as msg:
            raise SystemExit(msg)
        # except Exception as err:
        #     raise err


import Tkinter

Tkinter.CallWrapper = TkErrorCatcher


def get_args():
    parser = argparse.ArgumentParser()
    # feature information
    parser.add_argument('--gui', action='store_true', help='GUI를 실행합니다.'.decode('utf-8'))
    parser.add_argument('--sa8', action='store_true', help='SA9파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa9', action='store_true', help='SA9파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa9_1', action='store_true', help='SA9-1파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa10', action='store_true', help='SA10파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa10_1', action='store_true', help='SA10-1파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa11', action='store_true', help='SA11파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa12', action='store_true', help='SA12파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa13', action='store_true', help='SA13파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa14', action='store_true', help='SA14파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--sa15', action='store_true', help='SA15파일을 변환합니다.'.decode('utf-8'))
    parser.add_argument('--path', type=str, default=os.getcwd(),
                        help='엑셀파일이 들어있는 폴더의 경로를 입력하세요.'.decode('utf-8'))
    parser.add_argument('--save_path', type=str, default=os.getcwd() + '/pdf',
                        help='PDF를 저장할 폴더의 경로를 입력하세요.'.decode('utf-8'))
    parser.add_argument('--filename', type=str, default='PDF', help='엑셀파일 이름을 입력하세요.'.decode('utf-8'))
    parser.add_argument('--title', type=str, default='', help='시험제목을 입력하세요.'.decode('utf-8'))
    parser.add_argument('--description', default='', type=str, help='시험 설명을 입력하세요.'.decode('utf-8'))
    return parser.parse_args()


# document에 들어가는 글에 스타일을 지정할 수 있다.
def style(document):
    style_1 = document.styles.add_style('Heading_1', docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
    style_1.base_style = document.styles['Heading 1']
    # style_1.font.name = 'HY 견고딕'.decode('utf-8')
    # style_1._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY 견고딕'.decode('utf-8'))  # 한글 폰트를 따로 설정해 준다
    style_1.font.size = docx.shared.Pt(15)
    style_1.font.bold = True
    style_1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    style_2 = document.styles.add_style('text', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    style_2.font.name = '돋움체'.decode('utf-8')                                   # 글씨체
    style_2._element.rPr.rFonts.set(qn('w:eastAsia'), '돋움체'.decode('utf-8'))    # 한글 폰트를 따로 설정해 준다
    style_2.font.size = docx.shared.Pt(10)
    style_2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return style_1, style_2


def _get_xlrd_cell_value(cell):
    value = cell.value
    if cell.ctype == xlrd.XL_CELL_DATE:
        datetime_tup = xlrd.xldate_as_tuple(value, 0)
        if datetime_tup[0:3] == (0, 0, 0):  # time format without date
            value = datetime.time(*datetime_tup[3:])
        else:
            value = datetime.datetime(*datetime_tup)
    return value


# xls파일을 xlsx로 변환
def xls2xlsx(name, path, **kw):
    xls_name = str(str(path) + '\\' + str(name) + ".xls")
    book_xls = xlrd.open_workbook(xls_name, formatting_info=True, ragged_rows=True, **kw)
    book_xlsx = openpyxl.workbook.Workbook()

    sheet_names = book_xls.sheet_names()
    for sheet_index in range(len(sheet_names)):
        sheet_xls = book_xls.sheet_by_name(sheet_names[sheet_index])
        if sheet_index == 0:
            sheet_xlsx = book_xlsx.active
            sheet_xlsx.title = sheet_names[sheet_index]
        else:
            sheet_xlsx = book_xlsx.create_sheet(title=sheet_names[sheet_index])
        for c_range in sheet_xls.merged_cells:
            rlo, rhi, clo, chi = c_range
            sheet_xlsx.merge_cells(start_row=rlo + 1, end_row=rhi,
                                   start_column=clo + 1, end_column=chi, )
        for row in range(sheet_xls.nrows):
            sheet_xlsx.append((
                _get_xlrd_cell_value(cell)
                for cell in sheet_xls.row_slice(row, end_colx=sheet_xls.row_len(row))
            ))
        book_xlsx.save(path + '/' + name + '.xlsx')


# def change_font():
#     path = 'C:\\WINDOWS\\Fonts\\NanumBarunGothic.ttf'
#     fontprop = fm.FontProperties(fname=path, size=18)
#     return fontprop


def load_excel(name, dic_path):
    path = str(str(dic_path) + '\\' + str(name) + ".xlsx")
    try:
        load_wb = load_workbook(path)
    except IOError:  # xls파일인 경우
        xls2xlsx(name, dic_path)
        load_wb = load_workbook(path)

    except TypeError as e:
        print "오류가 발생하였습니다. 파일을 다른이름으로 저장 후 다시 해보시기 바랍니다.".decode('utf-8')
        raise e
    return load_wb


def create_folder(directory):
    directory = directory
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print 'Error: Creating directory. ' + directory


def make_pdf(name, sa15_path, save_path):
    # 파일 경로 절대경로로
    in_file = os.path.abspath(sa15_path + '/doxs/' + name + '.docx')
    out_file = os.path.abspath(str(save_path) + '\\' + name)

    # word형식의 파일을 열기
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)

    # PDF형식으로 저장
    try:
        doc.SaveAs(out_file, FileFormat=17)
    except IOError as ioe:
        doc.Close()
        word.Quit()
        raise ioe
    except COMError as ce:
        doc.Close()
        word.Quit()
        print "저장하려는 PDF파일이 열려있습니다. 닫고 다시 실행해주세요.".decode('utf-8')
        raise ce
    doc.Close()
    word.Quit()


def boundary(lst):
    result = []
    for i in lst:
        temp = (1 - abs(i) ** 2) ** 0.5
        if i > 0:
            result.append(temp)
        else:
            result.append(-temp)
    return result


def convert_sa8(sa8_name, sa8_title, sa8_description, sa8_path, sa8_save_path):
    print '-----------------SA8 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa8_name, sa8_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    load_ws2 = load_wb['result_summary.csv']
    value = []
    for row2 in load_ws2.rows:
        row_value2 = []
        for cell2 in row2:
            row_value2.append(cell2.value)
        value.append(row_value2)
    column = value[0]
    value = pd.DataFrame(value[1:])
    value.columns = column
    csv_name = value['Dataset File'].values
    table_data = []
    for sheet_name in csv_name:
        load_ws3 = load_wb[sheet_name]
        value2 = []
        for row1 in load_ws3.rows:
            row_value2 = []
            for cell in row1:
                row_value2.append(cell.value)
            value2.append(row_value2)
        column = value2[0]
        value2 = pd.DataFrame(value2[1:])
        value2.columns = column
        table_data.append(value2)
        fig = plt.figure(figsize=(10, 6))
        ax1 = fig.add_subplot(1, 1, 1)
        ax2 = ax1.twinx()
        # fontprop = change_font() / fontproperties=fontprop
        line1 = ax1.plot(value2['TIME'], value2['AC_IRMS_1'], color='b', label='AC_IRMS_1')
        line2 = ax2.plot(value2['TIME'], value2['MC'], color='r', label='MC')
        ax1.set_xlabel('Time(secs)', size=10)
        ax1.set_ylabel('Current(A)', size=10)
        ax1.set_xlim(0, )
        lines = line1 + line2
        labels = ['AC_IRMS_1', 'MC']
        plt.legend(lines, labels, loc=1)
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa8_path + '/img')  # 폴더가 존재하는지 확인하고 없으면 생성
        plt.savefig(sa8_path + '/img/' + str(sheet_name.split('.')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa8_title = sa8_title.encode('utf-8')
    sa8_description = sa8_description.encode('utf-8')
    # 제목
    style_1, style_2 = style(document)
    document.add_paragraph(sa8_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험 설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa8_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws["B2"].value).decode('utf-8'), style=style_2)
    # 기능시험 결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')

    # 테이블 작성
    for i in range(len(csv_name)):
        document.add_page_break()
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        table = document.add_table(rows=1, cols=6, style='Light Shading')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '합격여부'.decode('utf-8')
        hdr_cells[1].text = 'Test'.decode('utf-8')
        hdr_cells[2].text = 'Power(%)'.decode('utf-8')
        hdr_cells[3].text = 'P(%)'.decode('utf-8')
        hdr_cells[4].text = 'Q(%)'.decode('utf-8')
        hdr_cells[5].text = '트립시험(s)'.decode('utf-8')
        for a, b, c, d, e, f in value.loc[:, 'Result':'t_trip_meas'].values.tolist():
            row_cells = table.add_row().cells
            row_cells[0].text = str(a)
            row_cells[1].text = str(b)
            row_cells[2].text = str(c)
            row_cells[3].text = str(d)
            row_cells[4].text = str(e)
            row_cells[5].text = str(f)
        document.add_picture(sa8_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='List Bullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa8_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa8_path + '/doxs/' + sa8_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa8_name, sa8_path, sa8_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa9(sa09_name, sa09_title, sa09_description, sa09_path, sa09_save_path):
    print '-----------------SA9 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa09_name, sa09_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.endswith('.csv') == True][0][1:].reset_index(drop=True).values
    all_x_axis = []
    all_y_axis = []
    all_graph_title = []
    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        x_axis = []
        y_axis = []
        graph_title = ''
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        value = pd.DataFrame(value[1:])
        param = sheet_name.split('_')
        if param[2] == 'all':
            x_axis = value[0]
            y_axis = value[[1, 2, 25]]
            y_axis.columns = [0, 1, 2]
            graph_title = 'Voltage Ride-Through LV1 All Phase, ' + str(param[3].split('.')[0]) + '% Power'
        elif param[2] == 'p1':
            x_axis = value[0]
            y_axis = value[[1, 2, 25]]
            y_axis.columns = [0, 1, 2]
            graph_title = 'Voltage Ride-Through LV1 Phase A, ' + str(param[3].split('.')[0]) + '% Power'
        elif param[2] == 'p2':
            x_axis = value[0]
            y_axis = value[[8, 9, 25]]
            y_axis.columns = [0, 1, 2]
            graph_title = 'Voltage Ride-Through LV1 Phase B, ' + str(param[3].split('.')[0]) + '% Power'
        elif param[2] == 'p3':
            x_axis = value[0]
            y_axis = value[[15, 16, 25]]
            y_axis.columns = [0, 1, 2]
            graph_title = 'Voltage Ride-Through LV1 Phase C, ' + str(param[3].split('.')[0]) + '% Power'
        else:
            print('index에 있는 sheet 이름 중 all,pl,p2,p3 가 3번째에 들어가 있지 않습니다.')

        all_x_axis.append(x_axis)
        all_y_axis.append(y_axis)
        all_graph_title.append(graph_title)

    for i in range(len(all_x_axis)):
        fig = plt.figure()
        ax1 = fig.add_subplot(1, 1, 1)
        # fontprop = change_font()
        ax2 = ax1.twinx()
        line1 = ax1.plot(all_x_axis[i], all_y_axis[i][0], color='b', label='AC_VRMS_A')
        line2 = ax2.plot(all_x_axis[i], all_y_axis[i][1], color='r', label='AC_IRMS_A')
        line3 = ax2.plot(all_x_axis[i], all_y_axis[i][2], color='k', linestyle='--', label='AC_IRMS_PASS')
        ax1.set_xlabel('Time(secs)', size=10)
        ax1.set_ylabel('Voltage(V)', size=10)
        ax2.set_ylabel('Current(A)', size=10)

        lines = line1 + line2 + line3
        labels = ['AC_VRMS_A', 'AC_IRMS_A', 'AC_IRMS_PASS']
        plt.title(all_graph_title[i].decode('utf-8'), size=15)
        plt.legend(lines, labels, loc=3)
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa09_path + '/img')  # 폴더가 존재하는지 확인하고 없으면 생성
        plt.savefig(sa09_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png')
    # 사용하기 위한 변수 선언
    document = Document()
    try:
        sa09_title = sa09_title.encode('utf-8')
        sa09_description = sa09_description.encode('utf-8')

        # 제목
        style_1, style_2 = style(document)  # 스타일 설정
        document.add_paragraph(sa09_title.decode('utf-8'), style=style_1)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        # 시험설명
        document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(sa09_description.decode('utf-8'), style=style_2)
        # 판정기준
        document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(str(load_ws["B2"].value).decode('utf-8'), style=style_2)

        # 기능시험결과
        document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
        for i in range(len(csv_name)):
            document.add_page_break()
            try:
                mer_title = p_title[i]
            except IndexError:
                mer_title = ''
            document.add_paragraph(mer_title, style='ListNumber')
            document.add_picture(sa09_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
            caption = '<' + mer_title + '>'  # 캡션 달기
            document.add_paragraph(caption.decode('utf-8'), style=style_2)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
            try:
                # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
                temp = r_title[i]
                mer_title2 = '* 결과검토: ' + temp
                document.add_paragraph(mer_title2, style=style_2)
            except KeyError:
                pass
            except IndexError:
                pass
        # 기능시험 결과 요약
        document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)
    except Exception as e:
        create_folder(sa09_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
        document.save(sa09_path + '/doxs/' + sa09_name + '.docx')
        print(e)
    create_folder(sa09_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa09_path + '/doxs/' + sa09_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa09_name, sa09_path, sa09_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa9_1(sa09_1_name, sa09_1_title, sa09_1_description, sa09_1_path, sa09_1_save_path):
    print '-----------------SA9-1 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa09_1_name, sa09_1_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.endswith('.csv') == True][0][1:].reset_index(drop=True).values
    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        column = value[0]
        value = pd.DataFrame(value[1:])
        value.columns = column
        fig = plt.figure()
        ax1 = fig.add_subplot(1, 1, 1)
        ax2 = ax1.twinx()
        # fontprop = change_font()
        line1 = ax1.plot(value['TIME'], value['U1'], color='c', label='AC_V_A')
        line2 = ax2.plot(value['TIME'], value['I1'], color='r', alpha=0.5, label='AC_I_A')
        ax1.set_xlabel('Time(secs)', size=10)
        ax1.set_ylabel('Voltage(V)', size=10)
        ax2.set_ylabel('Current(A)', size=10)
        ax1.set_xlim(0, )
        lines = line1 + line2
        labels = ['AC_V_A', 'AC_I_A']
        plt.title('Voltage Ride-Through (Trip time, Waveform)'.decode('utf-8'), size=15)
        plt.legend(lines, labels, loc=3)
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa09_1_path + '/img')  # 폴더가 존재하는지 확인하고 없으면 생성
        plt.savefig(sa09_1_path + '/img/' + str(sheet_name.split('.')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    try:
        sa09_1_title = sa09_1_title.encode('utf-8')
        sa09_1_description = sa09_1_description.encode('utf-8')

        # 제목
        style_1, style_2 = style(document)  # 스타일 설정
        document.add_paragraph(sa09_1_title.decode('utf-8'), style=style_1)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        # 시험설명
        document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(sa09_1_description.decode('utf-8'), style=style_2)
        # 판정기준
        document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

        # 기능시험결과
        document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
        for i in range(len(csv_name)):
            document.add_page_break()
            try:
                mer_title = p_title[i]
            except IndexError:
                mer_title = ''
            document.add_paragraph(mer_title, style='ListNumber')
            document.add_picture(sa09_1_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
            caption = '<' + mer_title + '>'  # 캡션 달기
            document.add_paragraph(caption.decode('utf-8'), style=style_2)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
            try:
                # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
                temp = r_title[i]
                mer_title2 = '* 결과검토: ' + temp
                document.add_paragraph(mer_title2, style=style_2)
            except KeyError:
                pass
            except IndexError:
                pass
        # 기능시험 결과 요약
        document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
        document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)
    except Exception as e:
        create_folder(sa09_1_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
        document.save(sa09_1_path + '/doxs/' + sa09_1_name + '.docx')
        print(e)

    create_folder(sa09_1_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa09_1_path + '/doxs/' + sa09_1_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa09_1_name, sa09_1_path, sa09_1_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa10(sa10_name, sa10_title, sa10_description, sa10_path, sa10_save_path):
    print '-----------------SA10 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa10_name, sa10_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.contains('.csv') == True][0][1:].reset_index(drop=True).values
    all_x_axis = []
    all_y_axis = []
    all_graph_title = []
    power = []
    load_ws3 = load_wb['result_summary.csv']
    for row in load_ws3.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        power.append(row_value)
    power = pd.DataFrame(power[1:])[4]
    for sheet in range(len(csv_name)):
        load_ws2 = load_wb[csv_name[sheet]]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        value = pd.DataFrame(value[1:])
        x_axis = value[0]
        y_axis = value[[7, 2, 25]]
        y_axis.columns = [0, 1, 2]
        graph_title = 'Frequency Ride-Through LF1 (' + str(power[sheet]) + '% Power)'
        all_x_axis.append(x_axis)
        all_y_axis.append(y_axis)
        all_graph_title.append(graph_title)

    for i in range(len(all_x_axis)):
        fig = plt.figure(figsize=(12, 6))
        ax1 = fig.add_subplot(1, 1, 1)
        ax2 = ax1.twinx()
        # fontprop = change_font()
        line1 = ax1.plot(all_x_axis[i], all_y_axis[i][0], color='b', label='AC_FREQ_A')
        line2 = ax2.plot(all_x_axis[i], all_y_axis[i][1], color='r', label='AC_IRMS_A')
        line3 = ax2.plot(all_x_axis[i], all_y_axis[i][2], color='k', linestyle='--', label='AC_IRMS_PASS')
        ax1.set_xlabel('Time(secs)', size=10)
        ax1.set_ylabel('Frequency(Hz)', size=10)
        ax2.set_ylabel('Current(A)', size=10)
        lines = line1 + line2 + line3
        labels = ['AC_FREQ_A', 'AC_IRMS_A', 'AC_IRMS_PASS']
        plt.title(all_graph_title[i].decode('utf-8'), size=15)
        plt.legend(lines, labels, loc=3)
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa10_path + '/img/')  # 폴더가 존재하는지 확인하고 없으면 생성
        plt.savefig(sa10_path +'/img/' + str(csv_name[i].split('.')[0]) + '.png')
    # 사용하기 위한 변수 선언
    document = Document()
    sa10_title = sa10_title.encode('utf-8')
    sa10_description = sa10_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa10_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa10_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(csv_name)):
        document.add_page_break()
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa10_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa10_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa10_path + '/doxs/' + sa10_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa10_name, sa10_path, sa10_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa10_1(sa10_1_name, sa10_1_title, sa10_1_description, sa10_1_path, sa10_1_save_path):
    print '-----------------SA10-1 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa10_1_name, sa10_1_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.endswith('.csv') == True][0][1:].reset_index(drop=True).values
    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        column = value[0]
        value = pd.DataFrame(value[1:])
        value.columns = column

        fig = plt.figure(figsize=(10, 6))
        ax1 = fig.add_subplot(1, 1, 1)
        ax2 = ax1.twinx()

        line1 = ax1.plot(value['TIME'], value['I1'], color='c', label='AC_I_A', linewidth=0.2)
        line2 = ax2.plot(value['TIME'], value['Target F'], color='r', label='Target F')
        ax1.set_xlabel('Time(secs)', size=10)
        ax1.set_ylabel('Current(A)', size=10)
        ax2.set_ylabel('Frequency (Hz)', size=10)
        ax1.set_xlim(0, )
        lines = line1 + line2
        labels = ['AC_I_A', 'Target F']
        plt.title('Frequency Ride-Through LF2 (Trip time, Waveform)'.decode('utf-8'), size=15)
        plt.legend(lines, labels, loc=1)
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa10_1_path + '/img')  # 폴더가 존재하는지 확인하고 없으면 생성
        plt.savefig(sa10_1_path + '/img/' + str(sheet_name.split('.')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa10_1_title = sa10_1_title.encode('utf-8')
    sa10_1_description = sa10_1_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa10_1_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa10_1_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(csv_name)):
        document.add_page_break()
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa10_1_path + '/img/' + str(csv_name[i].split('.')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa10_1_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa10_1_path + '/doxs/' + sa10_1_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa10_1_name, sa10_1_path, sa10_1_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa11(sa11_name, sa11_title, sa11_description, sa11_path, sa11_save_path):
    print '-----------------SA11 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa11_name, sa11_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.contains('.csv') == True][0].reset_index(drop=True).values

    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        value = pd.DataFrame(value[1:])

        df = value[[0, 3]]
        df.columns = ['TIME', 'AC_P_1']
        load_ws3 = load_wb[str(sheet_name.split('.csv')[0]) + '_plot']
        value2 = []
        for row2 in load_ws3.rows:
            row_value2 = []
            for cell2 in row2:
                row_value2.append(cell2.value)
            value2.append(row_value2)
        value2 = pd.DataFrame(value2[1:])
        value2.columns = ['time_min', 'min', 'time_max', 'max']
        value2 = value2.astype({'time_min': np.float, 'min': np.float, 'time_max': np.float, 'max': np.float})

        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(1, 1, 1)
        ax.plot(df['TIME'], df['AC_P_1'], color='b')
        ax.plot(value2['time_min'], value2['min'], color='r', linestyle='--', label='min')
        ax.plot(value2['time_max'], value2['max'], color='r', linestyle='--', label='max')
        plt.xlabel('Time(secs)', size=10)
        plt.ylabel('Power(W)', size=10)
        plt.legend(['AC_P_1', 'min', 'max'])
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa11_path + '/img')
        plt.savefig(sa11_path + '/img/' + str(sheet_name.split('.csv')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa11_title = sa11_title.encode('utf-8')
    sa11_description = sa11_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa11_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬

    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa11_description.decode('utf-8'), style=style_2)

    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(csv_name)):
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa11_path + '/img/' + str(csv_name[i].split('.csv')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약

    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa11_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa11_path + '/doxs/' + sa11_name + '.docx')

    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa11_name, sa11_path, sa11_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa12(sa12_name, sa12_title, sa12_description, sa12_path, sa12_save_path):
    print '-----------------SA12 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa12_name, sa12_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1].dropna(axis=0).reset_index(drop=True)  # 테이블 제목
    r_title = all_value[2].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    d_column = [i for i in all_value[3].unique() if str(i) != 'nan']

    result_summary = str(load_ws['A2'].value)
    load_ws2 = load_wb[result_summary]

    all_values = []
    for row in load_ws2.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_values.append(row_value)
    column = all_values[0]
    data = all_values[1:]
    df = pd.DataFrame(data=data, columns=column)
    df2 = df[['Power Level (%)', 'Iteration', 'PF Target', 'PF Actual 1', 'PF Actual 2', 'PF Actual 3']]
    df3 = df[
        ['PF Target', 'Power 1 (pu)', 'Reactive Power 1 (pu)', 'Power 2 (pu)', 'Reactive Power 2 (pu)', 'Power 3 (pu)',
         'Reactive Power 3 (pu)', 'P Target at Rated (pu)', 'Q Target at Rated (pu)']]

    fig = plt.figure()
    ax = fig.add_subplot(1, 1, 1)
    # circle 그리기
    circle_center = (0, 0)
    circle_radius = 1
    for i in range(10):
        c = plt.Circle(circle_center, circle_radius, fc='w', ec='#A9A9A9', label='_circle axis')
        ax.add_patch(c)
        circle_radius -= 0.1
    # circle 보조축
    for i in range(0, 361, 30):
        radian = i * math.pi / 180
        start_x = 0 + (math.cos(radian) * 0.1)
        start_y = 0 + (math.sin(radian) * 0.1)
        end_x = 0 + (math.cos(radian) * 1)
        end_y = 0 + (math.sin(radian) * 1)
        plt.plot([start_x, end_x], [start_y, end_y], color='#A9A9A9', label='_circle axis')

    # 시험 결과 가져와서 나타내기
    target_lst = np.sort(df2['PF Target'].unique())
    c_lst = ['b', 'g', 'r', 'c', 'y']  # marker 색깔
    m_lst = ['v', 'o', 's']  # marker 모양
    x_lst = ['Power 1 (pu)', 'Power 2 (pu)', 'Power 3 (pu)']
    y_lst = ['Reactive Power 1 (pu)', 'Reactive Power 2 (pu)', 'Reactive Power 3 (pu)']
    pf_min = boundary(df['PF Min Allowed'].unique())  # PF Pass.Fail
    pf_max = boundary(df['PF Max Allowed'].unique())  # PF Pass.Fail

    data_label = ['Phase A, PF = ', 'Phase B, PF =', 'Phase C, PF =']

    for i in range(len(target_lst)):
        temp = df3[df3['PF Target'] == target_lst[i]]

        if i == 0:
            ax.plot([0, list(temp['P Target at Rated (pu)'])[0]], [0, list(temp['Q Target at Rated (pu)'])[0]]
                    , color='k', linestyle='--', label='PF target')
            ax.plot([0, (1-(pf_min[i] ** 2)) ** 0.5], [0, pf_min[i]], color='r', linestyle='--',
                    label='_PF Pass/Fail Boundary')
            ax.plot([0, (1-(pf_max[i] ** 2)) ** 0.5], [0, pf_max[i]], color='r', linestyle='--',
                    label='PF Pass/Fail Boundary')
        else:    # 두번째부터 범례 표시X
            ax.plot([0, list(temp['P Target at Rated (pu)'])[0]], [0, list(temp['Q Target at Rated (pu)'])[0]],
                    color='k', linestyle='--', label='_PF target')
            ax.plot([0, (1-(pf_min[i] ** 2)) ** 0.5], [0, pf_min[i]], color='r', linestyle='--',
                    label='_PF Pass/Fail Boundary')
            ax.plot([0, (1-(pf_max[i] ** 2)) ** 0.5], [0, pf_max[i]], color='r', linestyle='--',
                    label='_PF Pass/Fail Boundary')

        for j in range(len(x_lst)):
            ax.plot(temp[str(x_lst[j])], temp[str(y_lst[j])], linestyle='', marker=m_lst[j], c=c_lst[i], markersize=5,
                    alpha=0.5, label=(data_label[j] + str(target_lst[i])))

    ax.set_aspect('equal')  # 정사각형
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1)
    ax.set_xticks([-1.0, 0, +1.0])
    ax.set_xticklabels(["-100", "0", "100"])
    ax.set_yticks([-1.0, 0, +1.0])
    ax.set_yticklabels(["-100", "0", "100"])
    ax.set_xlabel('Active Power (% nameplate)', size=10)
    ax.set_ylabel('Reactive Power(% nameplate)', size=10)
    ax.set_title('Specified Power Factor', size=10)
    plt.legend(loc=3, fontsize='small')  # 범례 위치
    plt.grid(False)
    create_folder(sa12_path + '/img')
    plt.savefig(sa12_path + '/img/' + str(sa12_name) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa12_title = sa12_title.encode('utf-8')
    sa12_description = sa12_description.encode('utf-8')
    # 제목
    style_1, style_2 = style(document)
    document.add_paragraph(sa12_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험 설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa12_description.decode('utf-8'), style=style_2)

    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws["B2"].value).decode('utf-8'), style=style_2)

    # 기능시험 결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    document.add_picture(sa12_path + '/img/' + str(sa12_name) + '.png', width=Inches(5))  # 그림 불러와서 넣기
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    caption = '<' + str('Specified Power Factor 시험 전체 결과') + '>'  # 캡션 달기
    document.add_paragraph(caption.decode('utf-8'), style=style_2)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬

    num = [i for i in target_lst if i in d_column]
    # 테이블 작성
    document.add_page_break()
    for i in range(len(num)):
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        table = document.add_table(rows=1, cols=6, style='Light Shading')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '출력 (%)'.decode('utf-8')
        hdr_cells[1].text = '반복횟수'.decode('utf-8')
        hdr_cells[2].text = '목표 역률'.decode('utf-8')
        hdr_cells[3].text = '실제 역률 (A)'.decode('utf-8')
        hdr_cells[4].text = '실제 역률 (B)'.decode('utf-8')
        hdr_cells[5].text = '실제 역률 (C)'.decode('utf-8')

        for a, b, c, d, e, f in df2[df2['PF Target'] == num[i]].values.tolist():
            row_cells = table.add_row().cells
            row_cells[0].text = str(a)
            row_cells[1].text = str(b)
            row_cells[2].text = str(c)
            row_cells[3].text = str(d)
            row_cells[4].text = str(e)
            row_cells[5].text = str(f)
        try:
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='List Bullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa12_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa12_path + '/doxs/' + sa12_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa12_name, sa12_path, sa12_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa13(sa13_name, sa13_title, sa13_description, sa13_path, sa13_save_path):
    print '-----------------SA13 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa13_name, sa13_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    r_title = all_value[2].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    result_summary = str(load_ws['A2'].value)
    load_ws2 = load_wb[result_summary]
    all_values = []
    for row in load_ws2.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_values.append(row_value)
    column = all_values[0]
    data = all_values[1:]
    df = pd.DataFrame(data=data, columns=column)
    # 그림 만들기
    index_down = df[df['Dataset File'].str.contains('down') == True].index
    index_up = df[df['Dataset File'].str.contains('up') == True].index
    create_folder(sa13_path + '/img')  # 폴더가 존재하는지 확인하고 없으면 생성
    for i in range(len(index_down)):
        fig = plt.figure(figsize=(10, 6))
        img_title = df['Dataset File'][index_down[i]].decode("UTF-8").encode("EUC-KR")
        if i == 0:
            vv_1_1000 = df[0:index_down[i]]
        else:
            vv_1_1000 = df[index_up[i - 1] + 1:index_down[i]]
        ax = fig.add_subplot(1, 1, 1)
        ax.plot(vv_1_1000['Average Voltage (pu)'], vv_1_1000['Var Actual 1'] / 4000, linestyle='', marker='o',
                color='blue', label='Power')
        ax.plot(vv_1_1000['Average Voltage (pu)'], vv_1_1000['Var Target 1'] / 4000, color='black', label='VV curve')
        ax.plot(vv_1_1000['Average Voltage (pu)'], vv_1_1000['Var Min Allowed 1'] / 4000, linestyle=':', color='red',
                label='VV pass/fail band')
        ax.plot(vv_1_1000['Average Voltage (pu)'], vv_1_1000['Var Max Allowed 1'] / 4000, linestyle=':', color='red')
        ax.set_title('Volt-Var Function1', size=15)
        ax.set_xlabel('Grid Voltage(% nominal)', size=10)
        ax.set_ylabel('Reactive Power(% nameplate)', size=10)
        ax.set_xticks([0.9, 0.95, 1, 1.05, 1.1])
        ax.set_xticklabels(['90', '95', '100', '105', '110'])
        ax.set_yticks([-1.5, -1, -0.5, 0, 0.5, 1.0, 1.5])
        ax.set_yticklabels(['-150', '-100', '-50', '0', '50', '100', '150'])
        plt.savefig(sa13_path + '/img/' + img_title + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa13_title = sa13_title.encode('utf-8')
    sa13_description = sa13_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa13_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa13_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(index_down)):
        document.add_page_break()
        img_title = df['Dataset File'][index_down[i]]
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa13_path + '/img/' + str(img_title) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa13_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa13_path + '/doxs/' + sa13_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa13_name, sa13_path, sa13_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa14(sa14_name, sa14_title, sa14_description, sa14_path, sa14_save_path):
    print '-----------------SA14 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa14_name, sa14_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')

    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.contains('.csv') == True][0][1:].reset_index(drop=True).values
    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        column = value[0]
        value = pd.DataFrame(value[1:])
        value.columns = column
        df = value[['AC_FREQ_1', 'AC_P_1', 'AC_P_2', 'AC_P_3', 'freq_set']]
        lst = []
        for i in range(len(df) - 1):
            if df['freq_set'][i] != df['freq_set'][i + 1]:
                lst.append(i)
        graph_df = df.loc[lst, :]
        graph_df['Y'] = graph_df['AC_P_1'] + graph_df['AC_P_2'] + graph_df['AC_P_3']
        load_ws3 = load_wb[str(sheet_name.split('.csv')[0]) + '_plot']
        value2 = []
        for row2 in load_ws3.rows:

            row_value2 = []
            for cell2 in row2:
                row_value2.append(cell2.value)
            value2.append(row_value2)
        column2 = value2[0]
        value2 = pd.DataFrame(value2[1:])
        value2.columns = column2
        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(1, 1, 1)
        # fontprop = change_font()
        ax.plot(graph_df['AC_FREQ_1'], graph_df['Y'], color='b', linestyle='', marker='o')
        ax.plot(value2['freq'], value2['target'], color='k', label='target')
        ax.plot(value2['freq'], value2['min'], color='r', linestyle='--', label='min')
        ax.plot(value2['freq'], value2['max'], color='r', linestyle='--', label='max')
        ax.set_xlim(int(value2['freq'][0]), int(value2['freq'][len(value2['freq'])-1]))
        plt.xlabel('Frequence(Hz)', size=10)
        plt.ylabel('Active Power(W)', size=10)
        plt.title('FW Characterastic Curve 2'.decode('utf-8'), size=15)
        plt.legend(['100% Power', 'FW curve', 'FW curve min', 'FW curve max'])
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa14_path + '/img')
        plt.savefig(sa14_path + '/img/' + str(sheet_name.split('.csv')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa14_title = sa14_title.encode('utf-8')
    sa14_description = sa14_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa14_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa14_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(csv_name)):
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa14_path + '/img/' + str(csv_name[i].split('.csv')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa14_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa14_path + '/doxs/' + sa14_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa14_name, sa14_path, sa14_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


def convert_sa15(sa15_name, sa15_title, sa15_description, sa15_path, sa15_save_path):
    print '-----------------SA15 연구파일 pdf변환을 시작합니다.---------------------'.decode('utf-8')
    load_wb = load_excel(sa15_name, sa15_path)
    print '-----------------Excel File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    # 시트이름으로 불러오기
    load_ws = load_wb['Index']
    all_value = []
    for row in load_ws.rows:
        row_value = []
        for cell in row:
            row_value.append(cell.value)
        all_value.append(row_value)
    all_value = pd.DataFrame(all_value[1:])
    p_title = all_value[1][1:].dropna(axis=0).reset_index(drop=True)  # 기능 시험 결과
    r_title = all_value[2][1:].dropna(axis=0).reset_index(drop=True)  # 결과 검토
    csv_name = all_value[all_value[0].str.contains('.csv') == True][0][1:].reset_index(drop=True).values
    for sheet_name in csv_name:
        load_ws2 = load_wb[sheet_name]
        value = []
        for row1 in load_ws2.rows:
            row_value = []
            for cell in row1:
                row_value.append(cell.value)
            value.append(row_value)
        column = value[0]
        value = pd.DataFrame(value[1:])
        value.columns = column
        df = value[['AC_VRMS_1', 'AC_P_1', 'AC_P_2', 'AC_P_3', 'volt_set']]
        lst = []
        for i in range(len(df) - 1):
            if df['volt_set'][i] != df['volt_set'][i + 1]:
                lst.append(i)
        graph_df = df.loc[lst, :]
        graph_df['Y'] = graph_df['AC_P_1'] + graph_df['AC_P_2'] + graph_df['AC_P_3']
        plot_sheet = sheet_name.split('.csv')[0] + '_plot'
        load_ws3 = load_wb[str(plot_sheet.decode("UTF-8").encode("UTF-8"))]

        value2 = []
        for row2 in load_ws3.rows:
            row_value2 = []
            for cell2 in row2:
                row_value2.append(cell2.value)
            value2.append(row_value2)
        column2 = value2[0]
        value2 = pd.DataFrame(value2[1:])
        value2.columns = column2
        fig = plt.figure(figsize=(10, 6))
        ax = fig.add_subplot(1, 1, 1)
        ax.plot(graph_df['AC_VRMS_1'], graph_df['Y'], color='b', linestyle='', marker='o')
        ax.plot(value2['Voltage'], value2['target'], color='k', label='target')
        ax.plot(value2['Voltage'], value2['min'], color='r', linestyle='--', label='min')
        ax.plot(value2['Voltage'], value2['max'], color='r', linestyle='--', label='max')
        x_min = min(graph_df['AC_VRMS_1'].min(), value2['Voltage'].min()) - 1
        x_max = max(graph_df['AC_VRMS_1'].max(), value2['Voltage'].max()) + 1
        ax.set_xlim(x_min, x_max)
        plt.xlabel('Frequence(Hz)', size=10)
        plt.ylabel('Active Power(W)', size=10)
        plt.legend(['100% Power', 'FW curve', 'FW curve min', 'FW curve max'])
        plt.grid(True)
        fig.tight_layout()
        create_folder(sa15_path + '/img')
        plt.savefig(sa15_path + '/img/' + str(sheet_name.split('.csv')[0]) + '.png')

    # 사용하기 위한 변수 선언
    document = Document()
    sa15_title = sa15_title.encode('utf-8')
    sa15_description = sa15_description.encode('utf-8')

    # 제목
    style_1, style_2 = style(document)  # 스타일 설정
    document.add_paragraph(sa15_title.decode('utf-8'), style=style_1)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
    # 시험설명
    document.add_paragraph('시험 설명'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(sa15_description.decode('utf-8'), style=style_2)
    # 판정기준
    document.add_paragraph('판정기준'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['B2'].value).decode('utf-8'), style=style_2)

    # 기능시험결과
    document.add_paragraph('기능시험 결과'.decode('utf-8'), style='ListBullet')
    for i in range(len(csv_name)):
        try:
            mer_title = p_title[i]
        except IndexError:
            mer_title = ''
        document.add_paragraph(mer_title, style='ListNumber')
        document.add_picture(sa15_path + '/img/' + str(csv_name[i].split('.csv')[0]) + '.png', width=Inches(5))  # 그림 불러와서 넣기
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        caption = '<' + mer_title + '>'  # 캡션 달기
        document.add_paragraph(caption.decode('utf-8'), style=style_2)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # 중앙정렬
        try:
            # 결과검토 / 결과검토가 없을경우 발생하는 에러를 위해 try except구문
            temp = r_title[i]
            mer_title2 = '* 결과검토: ' + temp
            document.add_paragraph(mer_title2, style=style_2)
        except KeyError:
            pass
        except IndexError:
            pass
    # 기능시험 결과 요약
    document.add_paragraph('기능시험 결과 요약'.decode('utf-8'), style='ListBullet')
    document.add_paragraph(str(load_ws['C2'].value).decode('utf-8'), style=style_2)

    create_folder(sa15_path + '/doxs')  # 폴더가 존재하는지 확인하고 없으면 생성
    # docx파일을 생성을 위한 save('파일명')
    document.save(sa15_path + '/doxs/' + sa15_name + '.docx')
    print '-----------------Docs File을 성공적으로 불러왔습니다.---------------------'.decode('utf-8')
    make_pdf(sa15_name, sa15_path, sa15_save_path)
    print '-----------------PDF File을 성공적으로 만들었습니다.---------------------'.decode('utf-8')


if __name__ == '__main__':
    args = get_args()
    # -------------------------------------------------------------------------------------------------------------- #
    # command :
    #       - python excel2pdf.py --파일종류 --filename 파일이름 --title "시험제목" --description "시험설명"
    #       - python excel2pdf.py --파일종류 --filename 파일이름 --title "시험제목" --description "시험설명" --path 경로
    #       [파일이 다른 경로에 있을때]
    #
    #       ex) python excel2pdf.py --sa12 --filename SA12 --title "시험제목" --description "시험설명"
    #       ex) python excel2pdf.py --sa13 --filename SA13 --title "시험제목" --description "시험설명" --path C:\\python
    #
    # sa08/sa09/sa9_1/sa10/sa10_1/sa11/sa12/sa13/sa14/sa15 : 시험파일의 종류
    # filename : 해당 폴더에 있는 파일 이름
    # title : PDF 제목으로 들어가게될 내용
    # optional
    # description : PDF 내 시험설명으로 들어가게될 내용
    # path : 해당 엑셀파일이 있는 폴더의 경로
    # save_path : PDF의 저장 경로
    # -------------------------------------------------------------------------------------------------------------- #
    while args.sa8 or args.sa9 or args.sa9_1 or args.sa10 or args.sa10_1 or args.sa11 or args.sa12 or args.sa13 \
            or args.sa14 or args.sa15:
        args.title = unicode(args.title.decode('cp949'))
        args.description = unicode(args.description.decode('cp949'))
        try:
            if args.sa8:
                convert_sa8(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa9:
                convert_sa9(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa9_1:
                convert_sa9_1(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa10:
                convert_sa10(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa10_1:
                convert_sa10_1(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa11:
                convert_sa11(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa12:
                convert_sa12(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa13:
                convert_sa13(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa14:
                convert_sa14(args.filename, args.title, args.description, args.path, args.save_path)

            elif args.sa15:
                convert_sa15(args.filename, args.title, args.description, args.path, args.save_path)

        except IOError:
            print "파일을 찾지 못했습니다.".decode('utf-8')
            print 'path: ' + str(args.path).decode('utf-8')
            print 'file name: ' + args.filename.decode('utf-8')

        repeat = raw_input('Do you want to continue?(y/n)'.decode('utf-8'))
        if repeat == 'y' or repeat == 'Y':
            filename = raw_input('Enter the following file name.')
            title = raw_input('Enter the following title')
            description = raw_input('Enter the following file description')
            args.filename = filename
            args.title = title
            args.description = description
        else:
            break
    # -------------------------------------------------------------------------------------------------------------- #
    # command :
    #       - python excel2pdf.py --gui
    #
    # -------------------------------------------------------------------------------------------------------------- #
    if args.gui:
        root = excel2pdf_gui.init()
        excel2pdf_gui.Pdf(root)
        root.mainloop()
